import base64
import io
import logging
import re

from odoo import models, fields, api
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class KnowledgeDocument(models.Model):
    _name = "dc.knowledge.document"
    _description = "Dragon Coders Knowledge Base Document"
    _order = "create_date desc"

    name = fields.Char(string="Document Name", required=True)
    description = fields.Text(string="Description / Notes")
    file = fields.Binary(string="File", required=True, attachment=True)
    filename = fields.Char(string="Filename")
    file_type = fields.Selection(
        [
            ("pdf", "PDF"),
            ("docx", "Word Document"),
            ("txt", "Text File"),
            ("xlsx", "Excel Spreadsheet"),
        ],
        string="File Type",
        compute="_compute_file_type",
        store=True,
    )
    category = fields.Selection(
        [
            ("company", "Company Info"),
            ("services", "Services & Products"),
            ("projects", "Projects"),
            ("pricing", "Pricing"),
            ("technical", "Technical Docs"),
            ("faq", "FAQ"),
            ("other", "Other"),
        ],
        string="Category",
        default="other",
        required=True,
    )
    extracted_text = fields.Text(string="Extracted Text", readonly=True)
    chunk_count = fields.Integer(string="Chunks", compute="_compute_chunk_count")
    active = fields.Boolean(default=True)
    state = fields.Selection(
        [("pending", "Pending"), ("ready", "Ready"), ("error", "Error")],
        default="pending",
        string="Status",
    )
    error_msg = fields.Char(string="Error", readonly=True)

    @api.depends("filename")
    def _compute_file_type(self):
        for rec in self:
            fn = (rec.filename or "").lower()
            if fn.endswith(".pdf"):
                rec.file_type = "pdf"
            elif fn.endswith(".docx"):
                rec.file_type = "docx"
            elif fn.endswith(".txt"):
                rec.file_type = "txt"
            elif fn.endswith(".xlsx"):
                rec.file_type = "xlsx"
            else:
                rec.file_type = "txt"

    @api.depends("extracted_text")
    def _compute_chunk_count(self):
        for rec in self:
            rec.chunk_count = self.env["dc.knowledge.chunk"].search_count(
                [("document_id", "=", rec.id)]
            )

    def action_process(self):
        """Extract text from file and split into chunks."""
        for rec in self:
            try:
                text = rec._extract_text()
                if not text or not text.strip():
                    rec.state = "error"
                    rec.error_msg = "No text could be extracted from this file."
                    continue
                rec.extracted_text = text
                rec._create_chunks(text)
                rec.state = "ready"
                rec.error_msg = False
            except Exception as e:
                _logger.error("Knowledge doc error: %s", e)
                rec.state = "error"
                rec.error_msg = str(e)[:200]

    def _extract_text(self):
        """Extract plain text from the uploaded file."""
        file_data = base64.b64decode(self.file)
        ft = self.file_type

        if ft == "txt":
            return file_data.decode("utf-8", errors="ignore")

        elif ft == "pdf":
            try:
                import pdfplumber

                parts = []
                with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            parts.append(t)
                return "\n\n".join(parts)
            except ImportError:
                raise UserError("Run: pip install pdfplumber")

        elif ft == "docx":
            try:
                from docx import Document

                doc = Document(io.BytesIO(file_data))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                raise UserError("Run: pip install python-docx")

        elif ft == "xlsx":
            try:
                import openpyxl

                wb = openpyxl.load_workbook(
                    io.BytesIO(file_data), read_only=True, data_only=True
                )
                parts = []
                for sheet in wb.worksheets:
                    parts.append(f"[Sheet: {sheet.title}]")
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join(str(c) for c in row if c is not None)
                        if row_text.strip():
                            parts.append(row_text)
                return "\n".join(parts)
            except ImportError:
                raise UserError("Run: pip install openpyxl")

        return ""

    def _create_chunks(self, text, chunk_size=600, overlap=100):
        """Split text into overlapping chunks and store them."""
        # Delete old chunks first
        self.env["dc.knowledge.chunk"].search([("document_id", "=", self.id)]).unlink()

        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        words = text.split()
        chunks = []
        start = 0
        while start < len(words):
            chunk_text = " ".join(words[start : start + chunk_size]).strip()
            if chunk_text:
                chunks.append(
                    {
                        "document_id": self.id,
                        "content": chunk_text,
                        "category": self.category,
                        "sequence": len(chunks),
                    }
                )
            start += chunk_size - overlap

        if chunks:
            self.env["dc.knowledge.chunk"].create(chunks)

    def action_delete(self):
        """Delete document and its chunks."""
        self.env["dc.knowledge.chunk"].search(
            [("document_id", "in", self.ids)]
        ).unlink()
        self.unlink()
